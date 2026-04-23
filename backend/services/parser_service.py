"""
邮件解析 + MySQL 操作服务 — 从原 email_parser.py 迁移
MySQL 凭据改为从环境变量读取，不再硬编码
"""
from __future__ import annotations

import io
import os
import re
from datetime import datetime, date

from core.config import get_mysql_config

# ── 映射表（与 email_parser.py 保持一致） ────────────────────

PORT_MAP = {
    'iloilo': 'Iloilo', 'cebu': 'Cebu', 'davao': 'Davao',
    'cdo': 'CDO', 'cagayan': 'CDO',
    'mnl s': 'Manila South', 'mnl n': 'Manila North',
    'manila south': 'Manila South', 'manila north': 'Manila North',
    'manila': 'Manila South',
    '宿务': 'Cebu', '达沃': 'Davao', '卡加延': 'CDO',
    '马尼拉南': 'Manila South', '马尼拉北': 'Manila North',
    '马南': 'Manila South', '马北': 'Manila North', '伊洛伊洛': 'Iloilo',
}

SITG_PORT_CODE = {'DV': 'Davao', 'CB': 'Cebu', 'IO': 'Iloilo', 'CD': 'CDO', 'MN': 'Manila South'}

BL_PREFIX_LINE = {
    'SITTAG': 'SITC', 'SITG': 'SITC', 'OOLU': 'OOCL', 'COAU': 'COSCO',
    'MCLP': 'MCC', 'CNHU': 'CNC', 'CNH': 'CNC', 'EGLV': 'EVERGREEN',
    'HLCU': 'HAPAG-LLOYD', 'MAEU': 'MAERSK', 'MSCU': 'MSC', 'APHU': 'APL',
}

LINE_KEYWORDS = [
    ('WAN HAI', 'WAN HAI'), ('WANHAI', 'WAN HAI'), ('MAERSK', 'MAERSK'),
    ('COSCO', 'COSCO'), ('OOCL', 'OOCL'), ('SITC', 'SITC'),
    ('EMC', 'EMC'), ('CNC', 'CNC'), ('MCC', 'MCC'),
    ('EVERGREEN', 'EVERGREEN'), ('HAPAG', 'HAPAG-LLOYD'),
]

CONSIGNEE_BROKER = {
    'MAXWAY': 'Arden', 'POWERWAY': 'PGMC', 'GRANDWAY': 'PGMC',
    'UNIONBAY': 'Unionbay 施安定',
    'LFM': 'LFM', 'KRT CONSUMER GOODS': 'Arden', 'KRT': 'KRT', 'GBC': 'GBC',
}

FORWARDER_MAP = {
    '王斌': '王斌', '吉永': '吉永', '高阳': '高阳', '昊泽': '昊泽',
    '衍峰': '衍峰', '泓大': '泓大', '于丽英': '于丽英', '誉鼎': '誉鼎',
    '麒成': '麒成', '元一': '元一（工厂自理）', '海潮': '海潮（工厂自理）',
    '齐安': '齐安（工厂自理）',
}

FORWARDER_EMAIL_MAP = {'yongyue': '永越', 'xinfei': '鑫菲航', 'xinhang': '鑫菲航', 'opfs': 'OPFS'}

PRODUCT_EN = {
    'corrugated steel sheet': 'GL Corrugated Steel Sheet 彩涂瓦楞板',
    'corrugated': 'GL Corrugated Steel Sheet 彩涂瓦楞板',
    'welded wire mesh': 'Welded Wire Mesh 铁丝网',
    'welded wire': 'Welded Wire Mesh 铁丝网',
    'welding electrode': 'Welding Electrode 焊条',
    'cyclone wire': 'Cyclone Wire 铁围网',
    'barbed wire': 'Barbed Wire 刺绳',
    'gi wire': 'GI Wire 铁丝', 'steel wire': 'GI Wire 铁丝', 'tie wire': 'GI Wire 铁丝',
    'hog wire': 'Hog Wire 牛栏网', 'hardware cloth': 'Hardware Cloth 铁窗网',
    'steel matting': 'Steel Matting 冲花片', 'phenolic board': 'Phenolic Board 膜板',
    'plywood': 'Plywood 夹板', 'eco-board': 'Eco-Board 生态板', 'eco board': 'Eco-Board 生态板',
    'plastic canvas': 'Plastic Canvas 布', 'plastic resin': 'Plastic Resin 塑料粒',
    'lldpe': 'Plastic Resin 塑料粒', 'ldpe': 'Plastic Resin 塑料粒',
    'concrete nail': 'Concrete Nails 水泥钉', 'umbrella nail': 'Umbrella Nails 雨伞钉',
    'jetmatic pump': 'Jetmatic Pump 抽水泵', 'clamp': 'Clamp 夹码',
    'angle bar': 'Angle Bar 角钢', 'plain bar': 'Plain Bar 圆钢',
    'square bar': 'Square Bar 方钢', 'channel bar': 'Channel Bar 槽钢',
    't bar': 'T Bar T型钢', 'steel tube': 'Steel Tube 钢管', 'steel pipe': 'Steel Pipe 圆管',
    'steel sheet': 'Steel Sheet 钢板', 'steel strip': 'Steel Strips 钢带',
    'galvanized square tube': 'Galvanized Square Tube 镀锌方管',
    'square rectangular steel': 'Square/Rectangular Steel Tube 方矩管',
    'metal purling': 'Metal Purling 钢檩', 'upvc door': 'UPVC Door 塑钢门',
    'steel shovel': 'Steel Shovel 铁铲',
}

PRODUCT_CN = {
    '铁丝': 'GI Wire 铁丝', '牛栏网': 'Hog Wire 牛栏网',
    '勾花网': 'Cyclone Wire 铁围网', '铁围网': 'Cyclone Wire 铁围网',
    '网片': 'Steel Matting 冲花片', '胶合板': 'Plywood 胶合板',
    '夹板': 'Plywood 夹板', '膜板': 'Phenolic Board 膜板',
    '塑料米': 'Plastic Resin 塑料粒', '塑料粒': 'Plastic Resin 塑料粒',
    '钢板': 'Steel Sheet 钢板', '方管': 'Square Bar 方钢',
    '圆管': 'Steel Pipe 圆管', '钢管': 'Steel Tube 钢管',
    '钢带': 'Steel Strips 钢带', '角铁': 'Angle Bar 角钢',
    '角钢': 'Angle Bar 角钢', '龙骨': 'Metal Purling 钢檩',
    '钢檩': 'Metal Purling 钢檩', '刺丝': 'Barbed Wire 刺绳',
    '刺绳': 'Barbed Wire 刺绳', '水泥钉': 'Concrete Nails 水泥钉',
    '铁铲': 'Steel Shovel 铁铲', '彩板': 'GL Corrugated Steel Sheet 彩涂瓦楞板',
    '焊条': 'Welding Electrode 焊条',
}

FACTORY_NAMES = sorted([
    '远尚祥瑞', '君亦德', '森乐邦', '张荣迁', '盈美居', '高领域',
    '永伟', '元一', '东尚', '永达', '君乐', '华思', '力成', '汇隆',
    '捷鹏', '齐安', '昀翼', '山梅', '万通', '恒博', '华金', '璞丰',
    '德弘', '雄兴', '新源', '德蕴', '金星', '悦途', '恺丰', '前进',
    '天诺', '鼎丰', '鸿林', '永联', '莲顺', '浩森', '三发', '盈航',
    '东灏', '台正', '其他', '九维', '恒凯', '港众', '仁盛',
], key=len, reverse=True)

FIELD_TO_DB_COLUMN = {
    'B/L No.': 'B_L_No', 'Container No.': 'Container_No', 'POD': 'POD',
    '清关公司': '清关公司', 'Shipping Line': 'Shipping_Line',
    '实际清关费/柜': '实际清关费_柜', '货代': '货代', '合同号': '合同号',
    '合同号代码': '合同号代码', '订货单位': '订货单位',
    'Shipper': 'Shipper', 'Consignee': 'Consignee', 'Factory': 'Factory',
    '品名': '品名', '品名补充': '品名补充', '集装箱规格': '集装箱规格',
    '集装箱数量': '集装箱数量', 'Free Demurage': 'Free_Demurage',
    'Free Detention': 'Free_Detention', 'ETD': 'ETD', 'ETA': 'ETA',
    'Days to ETA': 'Days_to_ETA', '资料进度': '资料进度',
}

DB_TABLE = "Database_2026"


# ── MySQL 连接 ────────────────────────────────────────────────

def _get_mysql_conn():
    try:
        import pymysql
    except ImportError as exc:
        raise ImportError("请先安装 pymysql") from exc

    cfg = get_mysql_config()
    return pymysql.connect(
        host=cfg["host"],
        port=cfg["port"],
        user=cfg["user"],
        password=cfg["password"],
        database=cfg["database"],
        charset='utf8mb4',
        autocommit=False,
        ssl={'ssl_disabled': False},
    )


# ── 货代邮箱精确映射缓存 ──────────────────────────────────────
_forwarder_email_lookup_cache = None


def _load_forwarder_email_lookup_from_db():
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT `email_address`, `货代` FROM `forwarder_email_lookup`")
            rows = cur.fetchall()
        conn.close()
        return {email.lower(): forwarder for email, forwarder in rows}
    except Exception:
        return {}


def get_forwarder_email_lookup():
    global _forwarder_email_lookup_cache
    if _forwarder_email_lookup_cache is None:
        _forwarder_email_lookup_cache = _load_forwarder_email_lookup_from_db()
    return _forwarder_email_lookup_cache


def reload_forwarder_email_lookup():
    global _forwarder_email_lookup_cache
    _forwarder_email_lookup_cache = None


def init_forwarder_email_table():
    """建表并写入初始货代邮箱对照数据（一次性运行）"""
    initial_data = [
        ('caozuo02@tianjinyanfeng.com',     '衍峰'),
        ('liuying@yutongjieyun.com',         '王斌'),
        ('yaohong@yutongjieyun.com',         '王斌'),
        ('op6.xm@xmyuding.com',             '誉鼎'),
        ('opfs12@everbyd.com',              '永越'),
        ('donald.mai@everbyd.com',          '永越'),
        ('xmphi08@163.com',                 '鑫菲航'),
        ('zoey.zhu@hongdalogistics.com.cn', '泓大'),
        ('2379783890@qq.com',               '于丽英'),
        ('qdjy10@lslqd.com',                '吉永'),
    ]
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `forwarder_email_lookup` (
                  `id` INT AUTO_INCREMENT PRIMARY KEY,
                  `email_address` VARCHAR(200) NOT NULL UNIQUE,
                  `货代` VARCHAR(100) NOT NULL
                ) CHARACTER SET utf8mb4
            """)
            cur.executemany(
                "INSERT IGNORE INTO `forwarder_email_lookup` (`email_address`, `货代`) VALUES (%s, %s)",
                initial_data,
            )
            inserted = cur.rowcount
        conn.commit()
        print(f"forwarder_email_lookup 初始化完成，新增 {inserted} 条记录。")
    finally:
        conn.close()


# ── Consignee 精确对照缓存 ────────────────────────────────────
_consignee_lookup_cache = None


def _load_consignee_lookup_from_db():
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT `consignee_name`, `清关公司` FROM `consignee_lookup`")
            rows = cur.fetchall()
        conn.close()
        return {name.upper(): broker for name, broker in rows}
    except Exception:
        return {}


def get_consignee_lookup():
    global _consignee_lookup_cache
    if _consignee_lookup_cache is None:
        _consignee_lookup_cache = _load_consignee_lookup_from_db()
    return _consignee_lookup_cache


def reload_consignee_lookup():
    global _consignee_lookup_cache
    _consignee_lookup_cache = None


def init_consignee_table():
    """建表并写入初始 Consignee 对照数据（一次性运行）"""
    initial_data = [
        ('VESTA DRY GOODS TRADING', 'Gacutno'),
        ('BRIGHTLANE HOUSEHOLD SUPPLIES TRADING', 'PGMC'),
        ('EZELEA CONSUMER GOODS TRADING', 'PGMC'),
        ('LAO HOMELAND HARDWARE', 'PGMC'),
        ('JHI IMPORT EXPORT TRADING', 'PGMC'),
        ('SEEM INTERNATIONAL CORP.', 'Alin'),
        ('AC EIGHT CONSUMER GOODS TRADING', 'Arden'),
        ('BAP 888 CONSUMER GOODS TRADING', 'Arden'),
        ('KRT CONSUMER GOODS TRADING', 'Arden'),
        ('MARINOLD CONSUMER GOODS TRADING', 'Jerry'),
        ('MINASO DRY GOODS TRADING', 'Queo'),
        ('STUD CONSUMER GOODS TRADING', 'SDU'),
        ('VJK INTERNATIONAL IMPEX CORPORATION', 'Dahlia'),
        ('RODS SAGITTARIUS DRY GOODS TRADING', 'Emman'),
        ('PROVENIO MARKETING CORP.', 'Alin'),
        ('ALPHA STEELS ENTERPRISES CORPORATION', 'AP'),
        ('ARKKYL CONSUMER GOODS TRADING', 'Andrew'),
        ('SMALL GIANTS TRADING CORP.', 'Andrew'),
        ('LFM CONSUMER GOODS TRADING', 'Dash'),
        ('REVENUERO CONSUMER GOODS TRADING', 'Dash'),
        ('TIANTIAN GLOBAL FOOD PRODUCTS TRADING', 'Skylink'),
        ('FLC CONSUMER GOODS TRADING', 'Unionbay 施安定'),
        ('OEL LOGISTICS AND FORWARDING INC.', 'Unionbay 施安定'),
    ]
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `consignee_lookup` (
                  `id` INT AUTO_INCREMENT PRIMARY KEY,
                  `consignee_name` VARCHAR(200) NOT NULL UNIQUE,
                  `清关公司` VARCHAR(100) NOT NULL
                ) CHARACTER SET utf8mb4
            """)
            cur.executemany(
                "INSERT IGNORE INTO `consignee_lookup` (`consignee_name`, `清关公司`) VALUES (%s, %s)",
                initial_data,
            )
            inserted = cur.rowcount
        conn.commit()
        print(f"consignee_lookup 初始化完成，新增 {inserted} 条记录。")
    finally:
        conn.close()


# ── 清关公司邮箱 / 微信渠道对照缓存 ─────────────────────────────
_broker_email_lookup_cache = None
_wechat_broker_cache = None


def _load_broker_lookup_from_db():
    """从 broker_email_lookup 表加载，返回 (email_dict, wechat_set)。"""
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute(
                "SELECT `broker_name`, `email_address`, `channel` "
                "FROM `broker_email_lookup` ORDER BY `id`"
            )
            rows = cur.fetchall()
        conn.close()
        email_dict: dict[str, list[str]] = {}
        wechat_set: set[str] = set()
        for name, email, channel in rows:
            if channel == 'wechat':
                wechat_set.add(name)
            elif email:
                email_dict.setdefault(name, []).append(email)
            else:
                email_dict.setdefault(name, [])
        return email_dict, wechat_set
    except Exception:
        return {}, set()


def _ensure_broker_lookup():
    global _broker_email_lookup_cache, _wechat_broker_cache
    if _broker_email_lookup_cache is None:
        _broker_email_lookup_cache, _wechat_broker_cache = _load_broker_lookup_from_db()


def get_broker_emails() -> dict[str, list[str]]:
    _ensure_broker_lookup()
    return _broker_email_lookup_cache or {}


def get_wechat_brokers() -> set[str]:
    _ensure_broker_lookup()
    return _wechat_broker_cache or set()


def reload_broker_emails():
    global _broker_email_lookup_cache, _wechat_broker_cache
    _broker_email_lookup_cache = None
    _wechat_broker_cache = None


def init_broker_email_table():
    """建表并写入初始清关公司数据（一次性运行）"""
    initial_data = [
        ('Gacutno',        'jagacutno.lcb69@gmail.com',          'email'),
        ('PGMC',           'bonjaperson@yahoo.com',              'email'),
        ('PGMC',           'gbcb_srassoccha@yahoo.com',          'email'),
        ('PGMC',           'loistava_global@yahoo.com',          'email'),
        ('Arden',          'krttrading@yahoo.com',               'email'),
        ('Arden',          'jakeuberdon@yahoo.com',              'email'),
        ('Queo',           'wavephilip@yahoo.com',               'email'),
        ('Dahlia',         'kishera_angel05@yahoo.com',          'email'),
        ('Emman',          'emman_tagumpay@yahoo.com',           'email'),
        ('Andrew',         'ajbroker888@yahoo.com',              'email'),
        ('Dash',           'ling@dashcargologistics.com',        'email'),
        ('Dash',           'operations@dashcargologistics.com',  'email'),
        ('Goldrichline',   'goldrichline.rc@gmail.com',          'email'),
        ('Alin',           '',                                   'wechat'),
        ('Jerry',          '',                                   'wechat'),
        ('SDU',            '',                                   'wechat'),
        ('AP',             '',                                   'wechat'),
        ('Skylink',        '',                                   'wechat'),
        ('Unionbay 施安定', '',                                  'wechat'),
    ]
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `broker_email_lookup` (
                  `id` INT AUTO_INCREMENT PRIMARY KEY,
                  `broker_name` VARCHAR(100) NOT NULL,
                  `email_address` VARCHAR(255) NOT NULL DEFAULT '',
                  `channel` VARCHAR(20) NOT NULL DEFAULT 'email',
                  UNIQUE KEY `uq_broker_email` (`broker_name`, `email_address`)
                ) CHARACTER SET utf8mb4
            """)
            cur.executemany(
                "INSERT IGNORE INTO `broker_email_lookup` "
                "(`broker_name`, `email_address`, `channel`) VALUES (%s, %s, %s)",
                initial_data,
            )
            inserted = cur.rowcount
        conn.commit()
        print(f"broker_email_lookup 初始化完成，新增 {inserted} 条记录。")
    finally:
        conn.close()


# ── 产品名称对照缓存 ──────────────────────────────────────────
_product_name_lookup_cache = None


def _load_product_name_lookup_from_db():
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT `keyword`, `规范产品名` FROM `product_name_lookup`")
            rows = cur.fetchall()
        conn.close()
        return {kw: std for kw, std in rows}
    except Exception:
        return {}


def get_product_name_lookup():
    global _product_name_lookup_cache
    if _product_name_lookup_cache is None:
        _product_name_lookup_cache = _load_product_name_lookup_from_db()
    return _product_name_lookup_cache


def reload_product_name_lookup():
    global _product_name_lookup_cache
    _product_name_lookup_cache = None


def init_product_name_table():
    """建表并写入74条产品名称对照数据（来自Excel，一次性运行）"""
    initial_data = [
        ('CYCLONE WIRE',      'Cyclone Wire 勾花网'),
        ('勾花网',             'Cyclone Wire 勾花网'),
        ('PLYWOOD',           'Plywood 胶合板'),
        ('胶合板',             'Plywood 胶合板'),
        ('PHENOLIC BOARD',    'Phenolic Board 膜板'),
        ('膜板',               'Phenolic Board 膜板'),
        ('PVC BOARD',         'PVC Board'),
        ('PVC板',             'PVC Board'),
        ('SWIVEL CLAMP',      'Clamp 卡扣'),
        ('卡扣',               'Clamp 卡扣'),
        ('CONCRETE NAIL',     'Concrete Nails 水泥钉'),
        ('水泥钉',             'Concrete Nails 水泥钉'),
        ('UMBRELLA NAIL',     'Umbrella Nails 瓦楞钉'),
        ('瓦楞钉',             'Umbrella Nails 瓦楞钉'),
        ('JETMATIC PUMP',     'Jetmatic Pump 打水机'),
        ('打水机',             'Jetmatic Pump 打水机'),
        ('PLAIN BAR',         'Plain Bar 圆钢'),
        ('圆钢',               'Plain Bar 圆钢'),
        ('SQUARE BAR',        'Square Bar 方钢'),
        ('方钢',               'Square Bar 方钢'),
        ('HOG WIRE',          'Hog Wire 牛栏网'),
        ('牛栏网',             'Hog Wire 牛栏网'),
        ('PLASTIC CANVAS',    'Plastic Canvas 篷布'),
        ('篷布',               'Plastic Canvas 篷布'),
        ('CANVAS TARP',       'Plastic Canvas 篷布'),
        ('STEEL MATTING',     'Steel Matting 铁网片'),
        ('铁网片',             'Steel Matting 铁网片'),
        ('STEEL MAT',         'Steel Matting 铁网片'),
        ('WELDED WIRE MESH',  'Welded Wire Mesh 电焊网'),
        ('电焊网',             'Welded Wire Mesh 电焊网'),
        ('ANGLE BAR',         'Angle Bar 角钢'),
        ('角钢',               'Angle Bar 角钢'),
        ('ECO BOARD',         'Eco-Board 生态板'),
        ('生态板',             'Eco-Board 生态板'),
        ('WELDING ELECTRODE', 'Welding Electrode 电焊条'),
        ('电焊条',             'Welding Electrode 电焊条'),
        ('焊条',               'Welding Electrode 电焊条'),
        ('WELDING ROD',       'Welding Electrode 电焊条'),
        ('RESIN',             'Plastic Resin 塑料米'),
        ('PLASTIC RESIN',     'Plastic Resin 塑料米'),
        ('塑料米',             'Plastic Resin 塑料米'),
        ('T BAR',             'T Bar T型钢'),
        ('STEEL TUBE',        'Steel Tube 方管'),
        ('方管',               'Steel Tube 方管'),
        ('GI TUBE',           'Steel Tube 方管'),
        ('STEEL SHEET',       'Steel Sheet 钢板'),
        ('钢板',               'Steel Sheet 钢板'),
        ('彩板',               'Steel Sheet 钢板'),
        ('瓦楞板',             'Steel Sheet 钢板'),
        ('开平板',             'Steel Sheet 钢板'),
        ('STEEL STRIP',       'Steel Strip 带钢'),
        ('STEEL COIL',        'Steel Strip 带钢'),
        ('带钢',               'Steel Strip 带钢'),
        ('STEEL PIPE',        'Steel Pipe 圆管'),
        ('圆管',               'Steel Pipe 圆管'),
        ('GI PIPE',           'Steel Pipe 圆管'),
        ('UPVC DOOR',         'UPVC Door 塑料门'),
        ('CHANNEL BAR',       'Channel Bar 槽钢'),
        ('槽钢',               'Channel Bar 槽钢'),
        ('BARBED WIRE',       'Barbed Wire 刺绳'),
        ('刺绳',               'Barbed Wire 刺绳'),
        ('STEEL WIRE',        'Steel Wire 铁丝'),
        ('GI WIRE',           'Steel Wire 铁丝'),
        ('铁丝',               'Steel Wire 铁丝'),
        ('METAL PURLINS',     'Metal Purlins 龙骨'),
        ('龙骨',               'Metal Purlins 龙骨'),
        ('HARDWARE CLOTH',    'Hardware Cloth 方眼网'),
        ('方眼网',             'Hardware Cloth 方眼网'),
        ('STEEL SHOVEL',      'Steel Shovel 铁锹'),
        ('铁锹',               'Steel Shovel 铁锹'),
        ('PE FOAM',           'PE Foam'),
        ('FLAT BAR',          'Flat Bar 扁钢'),
        ('CLOUT NAIL',        'Clout Nail 油毡钉'),
        ('WALL CLIP',         'Wall Clip'),
    ]
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `product_name_lookup` (
                  `id` INT AUTO_INCREMENT PRIMARY KEY,
                  `keyword` VARCHAR(200) NOT NULL UNIQUE,
                  `规范产品名` VARCHAR(200) NOT NULL
                ) CHARACTER SET utf8mb4
            """)
            cur.executemany(
                "INSERT IGNORE INTO `product_name_lookup` (`keyword`, `规范产品名`) VALUES (%s, %s)",
                initial_data,
            )
            inserted = cur.rowcount
        conn.commit()
        print(f"product_name_lookup 初始化完成，新增 {inserted} 条记录。")
    finally:
        conn.close()


def _extract_forwarder_from_email(from_addr: str) -> str:
    """从发件人邮箱地址推断货代"""
    m = re.search(r'[\w.%+\-]+@[\w.\-]+', from_addr)
    email_only = m.group(0).lower() if m else from_addr.lower()
    email_lookup = get_forwarder_email_lookup()
    if email_only in email_lookup:
        return email_lookup[email_only]
    addr_lower = from_addr.lower()
    for keyword, forwarder in FORWARDER_EMAIL_MAP.items():
        if keyword in addr_lower:
            return forwarder
    return ''


# ── 解析工具函数 ──────────────────────────────────────────────

def _normalize_date(s: str) -> str:
    if not s:
        return ''
    s = s.strip()
    if re.match(r'^\d{4}-\d{2}-\d{2}$', s):
        return s
    m = re.match(r'^(\d{4})[/-](\d{1,2})[/-](\d{1,2})$', s)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3))).strftime('%Y-%m-%d')
        except ValueError:
            pass
    m = re.match(r'^(\d{1,2})\s*([A-Za-z]{3})\s*(\d{2,4})?$', s)
    if m:
        day, mon, yr_s = int(m.group(1)), m.group(2).capitalize(), m.group(3) or str(datetime.now().year)
        yr = int(yr_s) if len(yr_s) == 4 else 2000 + int(yr_s)
        mn = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,
              'Jul':7,'Aug':8,'Sep':9,'Oct':10,'Nov':11,'Dec':12}.get(mon)
        if mn:
            try:
                return date(yr, mn, day).strftime('%Y-%m-%d')
            except ValueError:
                pass
    return s


def _normalize_box(box_str: str) -> str:
    return box_str.upper().replace("'", "").replace("HQ", "HC")


def _bl_to_line(bl: str) -> str:
    bl = bl.upper()
    if re.match(r'^2\d{8}$', bl):
        return 'MCC'
    for prefix, line in BL_PREFIX_LINE.items():
        if bl.startswith(prefix):
            return line
    return ''


def _sitg_bl_to_port(bl: str) -> str:
    m = re.match(r'(SITG[A-Z]{2}|SITTAG)([A-Z]{2})', bl.upper())
    if m:
        return SITG_PORT_CODE.get(m.group(2), '')
    return ''


def _extract_port(text: str) -> str:
    for cn, std in PORT_MAP.items():
        if len(cn) > 2 and cn in text:
            return std
    for cn, std in PORT_MAP.items():
        if len(cn) == 2 and cn in text:
            return std
    text_up = text.upper()
    for key in sorted(PORT_MAP.keys(), key=len, reverse=True):
        if key in ('宿务','达沃','卡加延','马尼拉南','马尼拉北','马南','马北','伊洛伊洛'):
            continue
        if re.search(r'\b' + re.escape(key.upper()) + r'\b', text_up):
            return PORT_MAP[key]
    return ''


def _extract_products(text: str) -> str:
    found = []
    text_up = text.upper()
    product_lookup = get_product_name_lookup()
    if product_lookup:
        for kw in sorted(product_lookup.keys(), key=len, reverse=True):
            if kw.upper() in text_up:
                full = product_lookup[kw]
                if full not in found:
                    found.append(full)
    else:
        for kw in sorted(PRODUCT_EN.keys(), key=len, reverse=True):
            if kw.upper() in text_up:
                full = PRODUCT_EN[kw]
                if full not in found:
                    found.append(full)
        for cn, std in PRODUCT_CN.items():
            if cn in text and std not in found:
                found.append(std)
    return ','.join(found) if found else ''


def _extract_factories(text: str) -> str:
    matched, matched_set, covered = [], set(), []
    for name in FACTORY_NAMES:
        start = 0
        while True:
            idx = text.find(name, start)
            if idx == -1:
                break
            end = idx + len(name)
            if not any(s <= idx < e or s < end <= e for s, e in covered) and name not in matched_set:
                matched.append(name)
                matched_set.add(name)
                covered.append((idx, end))
            start = end
    matched.sort(key=lambda n: text.find(n))
    return ' '.join(matched) if matched else ''


# ── 核心解析函数 ──────────────────────────────────────────────

def parse_email_text(text: str) -> dict:
    """解析单段文本（标题/正文/附件），返回字段字典。"""
    result = {}
    text_up = text.upper()

    # 1. 提单号 — 主模式
    bl_patterns = [
        r'\b(2\d{8})\b',
        r'\b(SITTAG[A-Z]{2}\d{6,8})\b',
        r'\b(SITG[A-Z]{4}\d{6,8})\b',
        r'\b(OOLU[A-Z0-9]{8,12})\b',
        r'\b(COAU\d{10,13})\b',
        r'\b(MCLP[A-Z0-9]{8,12})\b',
        r'\b(CNHU[A-Z0-9]{8,12})\b',
        r'\b(EGLV[A-Z0-9]{8,12})\b',
        r'\b(HLCU[A-Z0-9]{8,12})\b',
        r'\b(MAEU[A-Z0-9]{8,12})\b',
        r'\b(CNH\d{7,10})\b',
        r'(?:提单号[：:]|提单[：:]?|B[/]?L[#\s]*[：:]?)\s*([A-Z0-9]{8,15})',
        r'\b([A-Z]{4}[A-Z0-9]{8,12})\b',
        r'\b(\d{12,13})\b',
    ]
    for pat in bl_patterns:
        m = re.search(pat, text_up)
        if m:
            bl = m.group(1)
            if re.match(r'^[A-Z]{4}\d{7}$', bl):
                continue
            result['B/L No.'] = bl
            break

    # 主模式未命中 → 兜底模式 + 排除词过滤
    if 'B/L No.' not in result:
        known_containers = set(re.findall(r'\b([A-Z]{4}\d{7})\b', text_up))
        exclude_words = set()
        exclude_words.update(k.upper() for k in PORT_MAP.keys())
        exclude_words.update(line.upper() for _, line in LINE_KEYWORDS)
        exclude_words.update({'GP', 'HC', 'HQ', 'ETD', 'ETA', 'BL', 'POD', 'POL'})
        for m in re.finditer(r'\b([A-Z]{1,6}[0-9][A-Z0-9]{5,13})\b', text_up):
            candidate = m.group(1)
            if re.match(r'^[A-Z]{4}\d{7}$', candidate):
                continue
            if candidate in known_containers:
                continue
            if re.match(r'^[A-Z]+$', candidate):
                continue
            if re.match(r'^\d{2}[A-Z]{2,3}$', candidate):
                continue
            if candidate in exclude_words:
                continue
            result['B/L No.'] = candidate
            result['_bl_fallback'] = True
            break

    # 2. 从 BL 推断船公司和港口
    if result.get('B/L No.'):
        bl = result['B/L No.']
        line = _bl_to_line(bl)
        if line:
            result['Shipping Line'] = line
        port = _sitg_bl_to_port(bl)
        if port:
            result['_port_from_bl'] = port

    # 3. 集装箱号
    containers = [c for c in re.findall(r'\b([A-Z]{4}\d{7})\b', text_up)
                  if c not in result.get('B/L No.', '')]
    if containers:
        result['Container No.'] = '\n'.join(containers)

    # 4. 集装箱规格和数量
    if '胶合板' in text:
        result['集装箱规格'] = '40HC'
        box_m = re.search(r'(\d+)\s*[xX×*＊]\s*(20GP|40H[CQ]|20\'?GP|40\'?H[CQ])', text_up)
        if box_m:
            result['集装箱数量'] = box_m.group(1)
        else:
            box_n = re.search(r'[（(](\d+)柜', text)
            if box_n:
                result['集装箱数量'] = box_n.group(1)
    else:
        box_m = re.search(r'(\d+)\s*[xX×*＊]\s*(20GP|40H[CQ]|20\'?GP|40\'?H[CQ])', text_up)
        if box_m:
            result['集装箱数量'] = box_m.group(1)
            result['集装箱规格'] = _normalize_box(box_m.group(2))
        else:
            size_cn = re.search(r'(\d+)(小|高)', text)
            if size_cn:
                result['集装箱数量'] = size_cn.group(1)
                result['集装箱规格'] = '20GP' if size_cn.group(2) == '小' else '40HC'
            else:
                box_type = re.search(r'\b(20GP|40H[CQ])\b', text_up)
                if box_type:
                    result['集装箱规格'] = _normalize_box(box_type.group(1))
                box_n = re.search(r'[（(](\d+)柜', text)
                if box_n and '集装箱数量' not in result:
                    result['集装箱数量'] = box_n.group(1)

    # 5. 目的港口
    port = _extract_port(text)
    if port:
        result['POD'] = port
    elif result.get('_port_from_bl'):
        result['POD'] = result['_port_from_bl']
    result.pop('_port_from_bl', None)

    # 6. 船公司 — 避免 SITG 开头 BL 误识别成 SITC
    if 'Shipping Line' not in result:
        for kw, line in LINE_KEYWORDS:
            if kw in text_up:
                if kw == 'SITC' and re.search(r'SITG[A-Z]', text_up) and 'SITC' not in text_up.split():
                    continue
                result['Shipping Line'] = line
                break

    # 7. 日期 — ETD（含船期、X.X计划兜底）、ETA
    etd_m = re.search(
        r'(?:ETD|装箱时间|开船时间)[：:\s/]*'
        r'(\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}\s*[A-Za-z]{3}\s*\d{0,4}|\d{1,2}[./]\d{1,2})',
        text
    )
    if etd_m:
        result['ETD'] = _normalize_date(etd_m.group(1))
    if 'ETD' not in result:
        cq_m = re.search(r'船期[：:\s]*(\d{1,2}[./]\d{1,2}|\d{4}[/-]\d{1,2}[/-]\d{1,2})', text)
        if cq_m:
            result['ETD'] = _normalize_date(cq_m.group(1))
    if 'ETD' not in result:
        plan_m = re.search(r'^(\d{1,2}[./]\d{1,2})计划', text.strip())
        if plan_m:
            result['ETD'] = _normalize_date(plan_m.group(1))

    eta_m = re.search(
        r'ETA[：:\s/]*'
        r'(\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}\s*[A-Za-z]{3}\s*\d{0,4}|\d{1,2}[./]\d{1,2})',
        text_up
    )
    if eta_m:
        result['ETA'] = _normalize_date(eta_m.group(1))

    # 8. 品名
    if '胶合板' in text:
        result['品名'] = 'Plywood 胶合板'
    else:
        products = _extract_products(text)
        if products:
            result['品名'] = products

    # 9. Consignee + 清关公司 — 先 DB 精确全称匹配（长度降序），后硬编码兜底
    lookup = get_consignee_lookup()
    for name_upper, broker in sorted(lookup.items(), key=lambda x: len(x[0]), reverse=True):
        if name_upper in text_up:
            result['Consignee'] = name_upper
            result['清关公司'] = broker
            result['_consignee_source'] = 'db'
            break
    if '清关公司' not in result:
        for key, broker in CONSIGNEE_BROKER.items():
            if key in text_up:
                result['清关公司'] = broker
                result['_consignee_source'] = 'fallback'
                break

    # 10. 货代（显示名中的中文关键词）
    for key, forwarder in FORWARDER_MAP.items():
        if key in text:
            result['货代'] = forwarder
            break

    # 11. 合同号
    contract_m = re.search(r'\b(\d{2}[A-Z]{3}\d{3,4})\b', text_up)
    if contract_m:
        result['合同号'] = contract_m.group(1)

    # 12. 工厂 — 先显式"工厂："前缀，再扫描关键词合并
    factory_m = re.search(r'工厂[：:]\s*(\S+)', text)
    if factory_m:
        explicit = factory_m.group(1)
        scanned = _extract_factories(text)
        if scanned and explicit not in scanned:
            result['Factory'] = explicit + ' ' + scanned
        else:
            result['Factory'] = scanned if scanned else explicit
    else:
        scanned = _extract_factories(text)
        if scanned:
            result['Factory'] = scanned

    # 13. 免柜期
    free_m = re.search(r'(\d+)\s*[+＋]\s*(\d+)', text)
    if free_m:
        result['Free Demurage'] = free_m.group(1)
        result['Free Detention'] = free_m.group(2)

    # 14. Shipper
    shipper_patterns = [
        r'SHIPPER[：:/\s]+([A-Z][A-Z0-9 .,&\'-]{3,60}?)(?:\n|$|CONSIGNEE|NOTIFY|ADDRESS)',
        r'发货人[：:\s]+(.{3,40}?)(?:\n|$)',
    ]
    for pat in shipper_patterns:
        shipper_m = re.search(pat, text_up if 'SHIPPER' in pat.upper() else text)
        if shipper_m:
            shipper_val = shipper_m.group(1).strip().rstrip('.,;')
            if len(shipper_val) >= 3:
                result['Shipper'] = shipper_val
            break

    # 15. Consignee 兜底（DB 未命中时从关键词/标题格式提取）
    if 'Consignee' not in result:
        consignee_m = re.search(
            r'-([A-Z][A-Z ]{4,50}(?:TRADING|INC\.?|CORP\.?|CO\.?|GOODS))\s*$',
            text_up
        )
        if consignee_m:
            result['Consignee'] = consignee_m.group(1).strip()
    if 'Consignee' not in result:
        cons_patterns = [
            r'CONSIGNEE[：:/\s]+([A-Z][A-Z0-9 .,&\'-]{3,60}?)(?:\n|$|NOTIFY|ADDRESS|PORT)',
            r'收货人[：:\s]+(.{3,40}?)(?:\n|$)',
        ]
        for pat in cons_patterns:
            cons_m = re.search(pat, text_up if 'CONSIGNEE' in pat.upper() else text)
            if cons_m:
                cons_val = cons_m.group(1).strip().rstrip('.,;')
                if len(cons_val) >= 3:
                    result['Consignee'] = cons_val
                break

    # 16. Days to ETA
    if result.get('ETA') and len(result['ETA']) == 10:
        try:
            eta_d = datetime.strptime(result['ETA'], '%Y-%m-%d').date()
            delta = (eta_d - date.today()).days
            result['Days to ETA'] = f"{abs(delta)}天前" if delta < 0 else f"{delta}天后"
        except ValueError:
            pass

    return result


def merge_results(base: dict, override: dict) -> dict:
    """合并两份解析结果：override 补齐 base 的空值。
    Consignee / 清关公司：DB 精确匹配可以覆盖硬编码兜底。
    """
    merged = dict(base)
    for k, v in override.items():
        if k.startswith('_'):
            continue
        if v and (not merged.get(k)):
            merged[k] = v
        elif k in ('Consignee', '清关公司') and v:
            if (override.get('_consignee_source') == 'db'
                    and merged.get('_consignee_source') == 'fallback'):
                merged[k] = v
    if (override.get('_consignee_source') == 'db'
            and merged.get('_consignee_source') == 'fallback'):
        merged['_consignee_source'] = 'db'
    return merged


# 向后兼容的旧别名
_merge = merge_results


def parse_full_email(subject: str, from_addr: str, body_text: str,
                     attachment_texts: list[dict] | None = None) -> dict:
    """综合解析：标题 → 正文 → 附件，标题识别到的 BL 始终优先。"""
    if attachment_texts is None:
        attachment_texts = []

    # 第一优先级：标题
    result = parse_email_text(subject or '')
    subject_bl = result.get('B/L No.', '') if not result.get('_bl_fallback') else ''

    # 第二优先级：正文
    if body_text:
        result = merge_results(result, parse_email_text(body_text))

    # 第三优先级：附件
    for att in attachment_texts:
        if att.get('text'):
            result = merge_results(result, parse_email_text(att['text']))

    # 标题若已明确识别 BL，覆盖正文中的手机号等噪音
    if subject_bl:
        result['B/L No.'] = subject_bl
        line = _bl_to_line(subject_bl)
        if line:
            result['Shipping Line'] = line
        port = _sitg_bl_to_port(subject_bl)
        if port:
            result['POD'] = port

    # 从发件人推断货代
    if not result.get('货代') and from_addr:
        for key, fw in FORWARDER_MAP.items():
            if key in from_addr:
                result['货代'] = fw
                break
        if not result.get('货代'):
            fw = _extract_forwarder_from_email(from_addr)
            if fw:
                result['货代'] = fw

    result['发件人'] = from_addr or ''
    result['邮件主题'] = subject or ''

    # 附件文件名列表
    if attachment_texts:
        filenames = [a.get('filename', '') for a in attachment_texts if a.get('filename')]
        if filenames:
            result['附件'] = ', '.join(filenames)

    # 清理内部标记
    result.pop('_bl_fallback', None)
    result.pop('_port_from_bl', None)
    result.pop('_consignee_source', None)
    return result


# ── 附件文字提取（服务端） ────────────────────────────────────

def extract_attachment_text(filename: str, data: bytes) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return _extract_pdf(data)
    elif fname.endswith(".docx"):
        return _extract_docx(data)
    elif fname.endswith((".xlsx", ".xlsm")):
        return _extract_excel_openpyxl(data)
    elif fname.endswith(".xls"):
        return _extract_excel_xlrd(data)
    return ""


def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text() or ""
            if t.strip():
                pages.append(t)
        return "\n\n".join(pages)
    except Exception as e:
        return ""


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(lines)
    except Exception:
        return ""


def _extract_excel_openpyxl(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None]
                if any(cells):
                    lines.append("  |  ".join(cells))
        wb.close()
        return "\n".join(lines)
    except Exception:
        return ""


def _extract_excel_xlrd(data: bytes) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        lines = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                cells = [str(sheet.cell_value(row_idx, c)).strip() for c in range(sheet.ncols)]
                if any(cells):
                    lines.append("  |  ".join(cells))
        return "\n".join(lines)
    except Exception:
        return ""


# ── 数据库操作 ────────────────────────────────────────────────

def insert_to_database(parsed: dict) -> dict:
    import time as _time

    columns, values = [], []
    for field_name, db_col in FIELD_TO_DB_COLUMN.items():
        val = parsed.get(field_name)
        if val is not None and val != '':
            if db_col in ('集装箱数量', 'Free_Demurage', 'Free_Detention'):
                try:
                    val = int(val)
                except (ValueError, TypeError):
                    pass
            columns.append(db_col)
            values.append(val)

    if '资料进度' not in columns:
        columns.append('资料进度')
        values.append('等待确认草稿')

    columns.append('创建时间')
    values.append(datetime.now().strftime('%Y-%m-%dT%H:%M:%S'))

    bl_no = parsed.get('B/L No.', '')
    col_str = ', '.join(f'`{c}`' for c in columns)
    placeholder_str = ', '.join(['%s'] * len(values))
    sql = f"INSERT INTO `{DB_TABLE}` ({col_str}) VALUES ({placeholder_str})"

    retry_keywords = ("timeout", "lost connection", "2013", "2006", "broken pipe")
    max_retries = 3
    last_exc: Exception | None = None

    for attempt in range(max_retries):
        conn = None
        try:
            conn = _get_mysql_conn()
            with conn.cursor() as cur:
                cur.execute(sql, values)
                row_id = cur.lastrowid
            conn.commit()
            return {"success": True, "row_id": row_id, "bl_no": bl_no,
                    "message": f"已写入数据库（id={row_id}，提单号={bl_no or '未识别'}）"}
        except Exception as e:
            last_exc = e
            if conn:
                try:
                    conn.rollback()
                except Exception:
                    pass
            err_lower = str(e).lower()
            if any(kw in err_lower for kw in retry_keywords) and attempt < max_retries - 1:
                _time.sleep(2)
                continue
            break
        finally:
            if conn:
                try:
                    conn.close()
                except Exception:
                    pass

    return {"success": False, "row_id": None, "bl_no": bl_no,
            "message": f"写入失败：{last_exc}"}


def update_progress_by_bl(bl_no: str, progress_value: str) -> dict:
    if not bl_no:
        return {"success": False, "rows": 0, "message": "提单号为空"}
    conn = None
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE `{DB_TABLE}` SET `资料进度` = %s WHERE `B_L_No` = %s",
                (progress_value, bl_no.upper())
            )
            rows = cur.rowcount
        conn.commit()
        return {"success": True, "rows": rows, "message": f"已更新 {rows} 条（{bl_no} → {progress_value}）"}
    except Exception as e:
        if conn:
            try:
                conn.rollback()
            except Exception:
                pass
        return {"success": False, "rows": 0, "message": f"更新失败：{e}"}
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


def update_mingxi_progress(bl_nos: list[str]) -> dict:
    if not bl_nos:
        return {"success": False, "updated": 0, "not_found": [], "message": "提单号列表为空"}
    conn = None
    try:
        conn = _get_mysql_conn()
        updated, not_found = 0, []
        with conn.cursor() as cur:
            for bl in bl_nos:
                cur.execute(
                    f"UPDATE `{DB_TABLE}` SET `明细_开单进度` = '已做明细' WHERE `B_L_No` = %s",
                    (bl.strip().upper(),)
                )
                (updated := updated + 1) if cur.rowcount else not_found.append(bl)
        conn.commit()
        msg = f"已更新 {updated} 条记录"
        if not_found:
            msg += f"，未找到：{', '.join(not_found)}"
        return {"success": True, "updated": updated, "not_found": not_found, "message": msg}
    except Exception as e:
        if conn:
            try:
                conn.rollback()
            except Exception:
                pass
        return {"success": False, "updated": 0, "not_found": [], "message": f"更新失败：{e}"}
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


def query_broker_by_bl(bl_no: str) -> str:
    if not bl_no:
        return ""
    conn = None
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute(
                f"SELECT `清关公司` FROM `{DB_TABLE}` WHERE `B_L_No` = %s ORDER BY `创建时间` DESC LIMIT 1",
                (bl_no.upper(),)
            )
            row = cur.fetchone()
            return row[0] if row and row[0] else ""
    except Exception:
        return ""
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


def query_product_by_bl(bl_no: str) -> str:
    """按提单号查询数据库中的品名，找不到返回空串。"""
    if not bl_no:
        return ""
    conn = None
    try:
        conn = _get_mysql_conn()
        with conn.cursor() as cur:
            cur.execute(
                f"SELECT `品名` FROM `{DB_TABLE}` WHERE `B_L_No` = %s ORDER BY `创建时间` DESC LIMIT 1",
                (bl_no.upper(),)
            )
            row = cur.fetchone()
            return str(row[0]).strip() if row and row[0] else ""
    except Exception:
        return ""
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


# ── 配置表操作（清关公司 / 草稿模板） ─────────────────────────

def _ensure_config_tables():
    """创建需要的配置表（模板 + broker 查找表）。"""
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `broker_email_lookup` (
                  `id` INT AUTO_INCREMENT PRIMARY KEY,
                  `broker_name` VARCHAR(100) NOT NULL,
                  `email_address` VARCHAR(255) NOT NULL DEFAULT '',
                  `channel` VARCHAR(20) NOT NULL DEFAULT 'email',
                  UNIQUE KEY `uq_broker_email` (`broker_name`, `email_address`)
                ) CHARACTER SET utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS `app_config_templates` (
                    `id` INT AUTO_INCREMENT PRIMARY KEY,
                    `name` VARCHAR(200) NOT NULL UNIQUE,
                    `body` TEXT NOT NULL,
                    `updated_at` TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) CHARSET=utf8mb4
            """)
        conn.commit()
    finally:
        conn.close()

    _migrate_legacy_brokers_if_needed()


def _migrate_legacy_brokers_if_needed():
    """把旧 app_config_brokers（JSON）数据一次性搬到 broker_email_lookup（仅当新表空）。"""
    import json
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            # 仅当旧表存在时才检查
            cur.execute(
                "SELECT COUNT(*) FROM information_schema.tables "
                "WHERE table_schema = DATABASE() AND table_name = 'app_config_brokers'"
            )
            if (cur.fetchone() or [0])[0] == 0:
                return
            cur.execute("SELECT COUNT(*) FROM `broker_email_lookup`")
            if (cur.fetchone() or [0])[0] > 0:
                return
            cur.execute("SELECT `name`, `emails_json` FROM `app_config_brokers`")
            legacy_rows = cur.fetchall()
            if not legacy_rows:
                return
            payload: list[tuple[str, str, str]] = []
            for name, emails_json in legacy_rows:
                try:
                    emails = json.loads(emails_json) if emails_json else []
                except Exception:
                    emails = []
                if not emails:
                    payload.append((name, '', 'email'))
                else:
                    for e in emails:
                        payload.append((name, str(e), 'email'))
            if payload:
                cur.executemany(
                    "INSERT IGNORE INTO `broker_email_lookup` "
                    "(`broker_name`, `email_address`, `channel`) VALUES (%s, %s, %s)",
                    payload,
                )
                conn.commit()
                print(f"[migration] 已从 app_config_brokers 迁移 {len(payload)} 条到 broker_email_lookup")
        reload_broker_emails()
    except Exception as e:
        print(f"[migration] 迁移 brokers 失败：{e}")
    finally:
        conn.close()


def get_brokers() -> dict:
    """返回 {name: {emails: [...], channel: 'email'|'wechat'}}。"""
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT `broker_name`, `email_address`, `channel` "
                "FROM `broker_email_lookup` ORDER BY `id`"
            )
            rows = cur.fetchall()
        out: dict[str, dict] = {}
        for name, email, channel in rows:
            entry = out.setdefault(name, {"emails": [], "channel": channel})
            # channel 以首次出现的为准，若存在 wechat 行则保持 wechat
            if channel == 'wechat':
                entry["channel"] = 'wechat'
            if email:
                entry["emails"].append(email)
        return out
    finally:
        conn.close()


def upsert_broker(name: str, emails: list[str], channel: str = "email") -> bool:
    """覆盖式写入：先删该 broker 的所有行，再重新插入当前 emails。"""
    channel = channel if channel in ("email", "wechat") else "email"
    if channel == "wechat":
        rows: list[tuple[str, str, str]] = [(name, '', 'wechat')]
    else:
        cleaned = [e.strip() for e in (emails or []) if e and e.strip()]
        rows = [(name, e, 'email') for e in cleaned] if cleaned else [(name, '', 'email')]

    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM `broker_email_lookup` WHERE `broker_name` = %s", (name,))
            cur.executemany(
                "INSERT INTO `broker_email_lookup` "
                "(`broker_name`, `email_address`, `channel`) VALUES (%s, %s, %s)",
                rows,
            )
        conn.commit()
        reload_broker_emails()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False
    finally:
        conn.close()


def delete_broker(name: str) -> bool:
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM `broker_email_lookup` WHERE `broker_name` = %s", (name,))
        conn.commit()
        reload_broker_emails()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False
    finally:
        conn.close()


def get_templates() -> list[dict]:
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT `name`, `body` FROM `app_config_templates` ORDER BY `id`")
            return [{"name": row[0], "body": row[1]} for row in cur.fetchall()]
    finally:
        conn.close()


def upsert_template(name: str, body: str) -> bool:
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO `app_config_templates` (`name`, `body`) VALUES (%s, %s) "
                "ON DUPLICATE KEY UPDATE `body` = VALUES(`body`)",
                (name, body)
            )
        conn.commit()
        return True
    except Exception:
        conn.rollback()
        return False
    finally:
        conn.close()


def delete_template(name: str) -> bool:
    conn = _get_mysql_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM `app_config_templates` WHERE `name` = %s", (name,))
        conn.commit()
        return True
    except Exception:
        conn.rollback()
        return False
    finally:
        conn.close()
